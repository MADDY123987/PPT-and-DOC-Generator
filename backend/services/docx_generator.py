from pathlib import Path
from typing import Dict, List
import re
import logging

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

# base storage dir (like you do for pptx)
BASE_DIR = Path(__file__).resolve().parent.parent
DOC_STORAGE_DIR = BASE_DIR / "storage" / "docs"
DOC_STORAGE_DIR.mkdir(parents=True, exist_ok=True)


def _clean_section_content(doc_title: str, heading: str, raw: str) -> str:
    """
    Post-process Gemini text so that:

    - Leading lines like:
        * "Electric Vehicle Market in India 2025" (doc title)
        * "Page 1 – Section 1", "Page 2 - Section 3", etc.
        * "Section 2: Growth Drivers"
        * a repeat of the heading, or "Heading: ..."
      are removed (and we keep stripping such lines until the
      first "normal" paragraph).
    - Literal '\\n' are converted to real newlines.
    - Extra empty lines at the start/end are removed.
    """
    if not raw:
        return ""

    # Normalise newlines
    text = raw.replace("\r\n", "\n").replace("\r", "\n")

    # Convert escaped '\n' from the model into real paragraph breaks
    text = text.replace("\\n", "\n")

    lines = [l.strip() for l in text.split("\n")]

    doc_title_low = (doc_title or "").strip().lower()
    heading_low = (heading or "").strip().lower()

    # Strip meta lines at the TOP until we hit a normal paragraph
    while lines:
        first = lines[0].strip()
        if not first:
            # blank line at top – just drop it
            lines.pop(0)
            continue

        f_low = first.lower()

        is_doc_title = f_low == doc_title_low
        is_page_section = f_low.startswith("page ") and "section" in f_low
        is_section_prefix = f_low.startswith("section ")
        is_heading_exact = f_low == heading_low
        is_heading_with_colon = heading_low and f_low.startswith(heading_low + ":")

        # Only strip the meta line if there *is* more content after it.
        # This prevents removing the only paragraph in cases where model output
        # matches the heading/title but there's no further text.
        if (is_doc_title or is_page_section or is_section_prefix or is_heading_exact or is_heading_with_colon):
            # If there's another non-empty line after this, pop current meta line
            if len(lines) > 1 and any(l.strip() for l in lines[1:]):
                lines.pop(0)
                continue
            else:
                # don't strip if it's the only content left
                break

        # First "normal" line reached – stop stripping
        break

    # Collapse duplicate blank lines
    cleaned: List[str] = []
    for line in lines:
        if not line:
            if cleaned and cleaned[-1] == "":
                continue
            cleaned.append("")
        else:
            cleaned.append(line)

    # Trim blank lines at start/end
    while cleaned and cleaned[0] == "":
        cleaned.pop(0)
    while cleaned and cleaned[-1] == "":
        cleaned.pop()

    return "\n".join(cleaned).strip()


# ----------------------------
# Helpers for page distribution
# ----------------------------

def split_section_into_parts(section: Dict[str, str], parts: int) -> List[Dict[str, str]]:
    """
    Split a single section's content into `parts` pieces by paragraph boundaries.
    Returns a list of section-like dicts with headings suffixed (only when needed).
    Each returned dict has keys: {"heading": ..., "content": ...}
    """
    heading = section.get("heading", "") or ""
    raw = section.get("content", "") or ""
    # Use existing cleaner to normalise before splitting
    content = _clean_section_content("", heading, raw)
    paragraphs = [p.strip() for p in content.split("\n") if p.strip()]

    if not paragraphs:
        # empty section: return empty parts
        return [{"heading": heading, "content": ""} for _ in range(parts)]

    # If paragraphs < parts, try split paragraphs into smaller chunks by sentences
    if len(paragraphs) < parts:
        # fallback: split text by approximate sentence count
        text = "\n".join(paragraphs)
        # Split sentences conservatively (keeps abbreviations mostly intact)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        if not sentences:
            # fallback to paragraphs joined
            sentences = [text]

        chunk_size = max(1, len(sentences) // parts)
        out = []
        idx = 0
        for i in range(parts):
            chunk = sentences[idx: idx + chunk_size]
            if not chunk:
                chunk = []
            out.append({
                "heading": f"{heading}" if i == 0 else f"{heading} (cont.)",
                "content": " ".join(chunk).strip()
            })
            idx += chunk_size
        # If leftover sentences, append to last
        leftover = sentences[idx:]
        if leftover:
            out[-1]["content"] = (out[-1]["content"] + " " + " ".join(leftover)).strip()
        return out

    # Otherwise distribute paragraphs evenly
    base = len(paragraphs) // parts
    rem = len(paragraphs) % parts
    out = []
    ptr = 0
    for i in range(parts):
        take = base + (1 if i < rem else 0)
        chunk_paras = paragraphs[ptr: ptr + take]
        ptr += take
        out.append({
            "heading": f"{heading}" if i == 0 else f"{heading} (cont.)",
            "content": "\n".join(chunk_paras).strip()
        })
    return out


def distribute_sections_across_pages(
    sections: List[Dict[str, str]],
    num_pages: int
) -> Dict[int, List[Dict[str, str]]]:
    """
    Distribute the flat list 'sections' into num_pages pages.

    Behavior:
    - If len(sections) >= num_pages: distribute sections as evenly as possible (some pages may have 1 extra).
    - If len(sections) < num_pages: split the longest section(s) into parts so each page has something to show.
    - Returns pages dict: {1: [...], 2: [...], ...}
    """
    if num_pages <= 0:
        num_pages = 1

    pages: Dict[int, List[Dict[str, str]]] = {i: [] for i in range(1, num_pages + 1)}
    total_sections = len(sections)

    if total_sections == 0:
        # nothing to distribute -> return empty pages
        return pages

    if total_sections >= num_pages:
        # Balanced distribution: base per page + remainder assigned to first pages
        base = total_sections // num_pages
        rem = total_sections % num_pages
        idx = 0
        for p in range(1, num_pages + 1):
            take = base + (1 if p <= rem else 0)
            pages[p] = sections[idx: idx + take]
            idx += take
        return pages

    # total_sections < num_pages -> we need to split some sections
    # strategy: find the longest sections by word count and split them to fill pages
    # compute word counts
    def word_count(s):
        return len(re.findall(r"\w+", s.get("content", "") or ""))

    # Sort sections descending by word count (longest first)
    sections_sorted = sorted(sections, key=word_count, reverse=True)

    # Create an array of page buckets (will convert to dict later)
    pages_list: List[List[Dict[str, str]]] = [[] for _ in range(num_pages)]

    # Place each original section into the first available pages (one per page up to number of sections)
    for i, sec in enumerate(sections_sorted):
        if i < num_pages:
            pages_list[i].append(sec)
        else:
            # should not happen because total_sections < num_pages here, but safe fallback:
            pages_list[0].append(sec)

    # Now fill empty pages by splitting the longest sections iteratively
    filled_pages = sum(1 for p in pages_list if p)
    empty_slots = num_pages - filled_pages

    # While there are empty slots, split the currently longest section available
    while empty_slots > 0:
        # find index of page with the longest total word count
        longest_page_idx = None
        longest_wc = 0
        for i, page in enumerate(pages_list):
            if not page:
                continue
            wc = sum(word_count(sec) for sec in page)
            if wc > longest_wc:
                longest_wc = wc
                longest_page_idx = i

        if longest_page_idx is None:
            # nothing left to split
            break

        # pop the first section from that page to split it
        sec_to_split = pages_list[longest_page_idx].pop(0)
        # Split it into 2 parts
        parts = split_section_into_parts(sec_to_split, 2)
        # Put first part back into original page, second part into next empty slot
        pages_list[longest_page_idx].append(parts[0])

        # find next empty slot
        next_empty = None
        for j, pg in enumerate(pages_list):
            if not pg:
                next_empty = j
                break
        if next_empty is None:
            # no empty slot found (shouldn't happen)
            break

        pages_list[next_empty].append(parts[1])
        empty_slots -= 1

    # Convert pages_list to dict with page numbers preserved in original order
    for idx, page in enumerate(pages_list):
        pages[idx + 1] = page

    return pages


# ----------------------------
# Original build_docx_file (unchanged aside from helpers above)
# ----------------------------

def build_docx_file(
    project_id: int,
    title: str,
    pages: Dict[int, List[Dict[str, str]]],  # {1: [...], 2: [...], ...}
) -> Path:
    """
    pages = {
      1: [
        {"heading": "Introduction", "content": "...."},
        {"heading": "Context", "content": "...."},
      ],
      2: [
        {"heading": "Market Analysis", "content": "...."},
      ],
      ...
    }

    - Each key is a page number (1-based).
    - Each value is a list of sections for that page.
    """

    # debug: show lengths of content per page
    try:
        page_debug = {p: [len(s.get("content", "")) for s in secs] for p, secs in pages.items()}
        logger.debug("Building DOCX for project %s — pages content lengths: %s", project_id, page_debug)
    except Exception:
        pass

    doc = Document()

    # Title page (Word will handle its own pagination)
    doc.add_heading(title, level=0)

    first_page = True

    for page_num in sorted(pages.keys()):
        # For the first *content* page after the title we don't add a break.
        # For every later content page we insert an explicit page break.
        if not first_page:
            doc.add_page_break()
        first_page = False

        page_sections = pages[page_num]
        num_sections = len(page_sections)

        # Simple font size logic based on how many sections in the page
        if num_sections <= 1:
            para_size = Pt(12)
            heading_size = Pt(16)
        elif num_sections == 2:
            para_size = Pt(11)
            heading_size = Pt(14)
        else:
            # 3 sections -> slightly smaller text
            para_size = Pt(10)
            heading_size = Pt(13)

        for section in page_sections:
            heading = section.get("heading", "") or ""
            raw_content = section.get("content", "") or ""

            # Clean up Gemini text (remove doc title / Page X / Section X lines, handle \n)
            content = _clean_section_content(title, heading, raw_content)

            if heading:
                h = doc.add_heading(heading, level=1)
                for run in h.runs:
                    run.font.size = heading_size

            if content:
                for para_text in content.split("\n"):
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        for run in p.runs:
                            run.font.size = para_size

    file_name = f"project_{project_id}.docx"
    file_path = DOC_STORAGE_DIR / file_name
    doc.save(file_path)

    return file_path
